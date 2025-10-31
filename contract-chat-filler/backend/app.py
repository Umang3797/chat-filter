from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
import os, tempfile
from placeholder_detect import extract_placeholders_from_text
from validation import validate_and_normalize
from docx_fill import fill_docx
from convert import doc_to_docx, docx_to_pdf
from session_store import InMemorySessions
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

SESSIONS = InMemorySessions()

@app.post("/api/session")
async def new_session():
    sid = SESSIONS.new()
    return {"session_id": sid}

@app.post("/api/upload")
async def upload(request: Request, file: UploadFile = File(...)):
    sid = request.headers.get("x-session-id")
    if not sid or not SESSIONS.get(sid):
        return JSONResponse({"error": "invalid session"}, status_code=400)

    suffix = os.path.splitext(file.filename or "")[1].lower()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(await file.read())
    tmp.close()

    src_path = tmp.name
    working_docx = None

    if suffix == ".docx":
        working_docx = src_path
    elif suffix == ".doc":
        working_docx = doc_to_docx(src_path)
    elif suffix == ".pdf":
        text = pdf_extract_text(src_path)
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        working_docx = tempfile.mktemp(suffix=".docx")
        doc.save(working_docx)
    else:
        return JSONResponse({"error": "Unsupported file type"}, status_code=400)

    doc = Document(working_docx)
    combined = []
    for p in doc.paragraphs:
        combined.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    combined.append(p.text)
    text = "\n".join(combined)

    placeholders = extract_placeholders_from_text(text)

    data = SESSIONS.get(sid)
    data.update({
        "working_docx": working_docx,
        "placeholders": placeholders,
        "values": {},
        "order": [p["key"] for p in placeholders],
        "cursor": 0
    })
    SESSIONS.set(sid, data)
    return {"placeholders": placeholders}

@app.get("/api/next")
async def next_question(request: Request):
    sid = request.headers.get("x-session-id")
    data = SESSIONS.get(sid)
    if not data:
        return JSONResponse({"error": "invalid session"}, status_code=400)
    order = data.get("order", [])
    cur = data.get("cursor", 0)
    while cur < len(order) and order[cur] in data["values"]:
        cur += 1
    data["cursor"] = cur
    SESSIONS.set(sid, data)
    if cur >= len(order):
        return {"done": True}
    key = order[cur]
    label = next((p["raw"] for p in data["placeholders"] if p["key"] == key), f"[{key}]")
    vtype = next((p["type"] for p in data["placeholders"] if p["key"] == key), "string")
    return {"done": False, "key": key, "label": label, "type": vtype}

@app.post("/api/answer")
async def submit_answer(request: Request, key: str = Form(...), value: str = Form(...)):
    sid = request.headers.get("x-session-id")
    data = SESSIONS.get(sid)
    if not data:
        return JSONResponse({"error": "invalid session"}, status_code=400)
    vtype = next((p["type"] for p in data["placeholders"] if p["key"] == key), "string")
    ok, norm, err = validate_and_normalize(vtype, value)
    if not ok:
        return {"ok": False, "error": err}
    data["values"][key] = norm
    data["cursor"] = min(len(data.get("order", [])), data.get("cursor", 0) + 1)
    SESSIONS.set(sid, data)
    return {"ok": True}

@app.post("/api/skip")
async def skip(request: Request):
    sid = request.headers.get("x-session-id")
    data = SESSIONS.get(sid)
    if not data:
        return JSONResponse({"error": "invalid session"}, status_code=400)
    data["cursor"] = min(len(data.get("order", [])), data.get("cursor", 0) + 1)
    SESSIONS.set(sid, data)
    return {"ok": True}

@app.post("/api/back")
async def back(request: Request):
    sid = request.headers.get("x-session-id")
    data = SESSIONS.get(sid)
    if not data:
        return JSONResponse({"error": "invalid session"}, status_code=400)
    data["cursor"] = max(0, data.get("cursor", 0) - 1)
    SESSIONS.set(sid, data)
    return {"ok": True}

@app.post("/api/fill")
async def fill(request: Request):
    sid = request.headers.get("x-session-id")
    data = SESSIONS.get(sid)
    if not data:
        return JSONResponse({"error": "invalid session"}, status_code=400)
    mapping = data.get("values", {})
    if not mapping:
        return JSONResponse({"error": "no values"}, status_code=400)

    filled_docx = tempfile.mktemp(suffix=".docx")
    fill_docx(data["working_docx"], filled_docx, mapping)

    data["filled_docx"] = filled_docx
    SESSIONS.set(sid, data)
    return {"ok": True}

@app.get("/api/download.docx")
async def download_docx(request: Request):
    sid = request.headers.get("x-session-id")
    data = SESSIONS.get(sid)
    if not data or not data.get("filled_docx"):
        return JSONResponse({"error": "not ready"}, status_code=400)
    return FileResponse(data["filled_docx"], media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="filled.docx")

@app.get("/api/preview.pdf")
async def preview_pdf(request: Request):
    sid = request.headers.get("x-session-id")
    data = SESSIONS.get(sid)
    if not data or not data.get("filled_docx"):
        return JSONResponse({"error": "not ready"}, status_code=400)
    pdf_path = docx_to_pdf(data["filled_docx"])
    return FileResponse(pdf_path, media_type="application/pdf", filename="filled.pdf")

# Serve the frontend (copied to /srv/site). Mount last so /api/* continues to work.
if os.path.isdir("/srv/site"):
    app.mount("/", StaticFiles(directory="/srv/site", html=True), name="frontend")
