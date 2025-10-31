from __future__ import annotations
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell
from typing import Dict, List
import re

PLACEHOLDER_PAT = re.compile(r"\[(?:(?!\[|\]).)*\]|\$?\[[_\s]+\]")

def replace_in_paragraph(paragraph: Paragraph, mapping: Dict[str, str]):
    if not paragraph.runs:
        return
    text = "".join(run.text for run in paragraph.runs)
    if not text:
        return

    def norm(s: str) -> str:
        s = s.strip()
        if s.startswith("[") and s.endswith("]"):
            s = s[1:-1]
        s = re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_")
        return s.lower()

    repls: List[tuple[int,int,str]] = []
    for m in PLACEHOLDER_PAT.finditer(text):
        raw = m.group(0)
        key = norm(raw)
        if key in mapping and mapping[key] is not None:
            repls.append((m.start(), m.end(), mapping[key]))

    if not repls:
        return

    pos = 0
    def append_text(chunk: str, style_src_idx: int):
        if chunk == "":
            return
        style_src = paragraph.runs[min(style_src_idx, len(paragraph.runs)-1)]
        r = paragraph.add_run()
        r.bold = style_src.bold
        r.italic = style_src.italic
        r.underline = style_src.underline
        r.style = style_src.style
        r.font.name = style_src.font.name
        r.font.size = style_src.font.size
        r.text = chunk

    # clear original
    for r in paragraph.runs: r.text = ""

    last_idx = 0
    for (start, end, value) in repls:
        before = text[pos:start]
        append_text(before, last_idx)
        append_text(value, last_idx)
        pos = end
        last_idx += 1
    append_text(text[pos:], last_idx)

def fill_docx(in_path: str, out_path: str, mapping: Dict[str, str]):
    doc = Document(in_path)

    def handle_block(container):
        for p in container.paragraphs:
            replace_in_paragraph(p, mapping)
        for t in container.tables:
            for row in t.rows:
                for cell in row.cells:
                    handle_block(cell)

    handle_block(doc)

    for section in doc.sections:
        if section.header:
            handle_block(section.header)
        if section.footer:
            handle_block(section.footer)

    doc.save(out_path)
